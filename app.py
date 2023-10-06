# import streamlit as st
# import os
# from docx import Document
# import re
# import openai
# import time
# import subprocess
# from graphviz import Digraph
# from PIL import Image

# # Function to extract text from a single Word file
# def extract_text_from_single_word_file(file):
#     doc = Document(file)
#     text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
#     clean_text = re.sub(r'[^\w\s]', '', text)

#     if len(clean_text) > 409050:
#         clean_text = clean_text[:400000]

#     extracted_text = clean_text[:40956]
#     return extracted_text

# # Function to create an interactive flowchart from the generated description
# def create_interactive_flowchart(description):
#     dot = Digraph(comment='Interactive Flowchart')
#     dot.attr(rankdir='LR')
#     dot.node('Start', 'Start', shape='ellipse', color='lightblue', style='filled')

#     lines = description.strip().split('\n')
#     for line in lines:
#         parts = line.strip().split('->')
#         if len(parts) == 2:
#             src, dest = parts
#             dot.edge(src.strip(), dest.strip())

#     # Save the flowchart as a PNG image
#     dot.render('flowchart', format='png')

#     # Display the flowchart image
#     return 'flowchart.png'

# # Initialize OpenAI API
# openai.api_key = 'sk-'  # Replace with your actual OpenAI API key

# # Streamlit UI
# st.title("Interactive Flowchart Generation from Text")

# # File upload
# st.header("Upload DOCX File")
# uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"])

# # Prompt input
# st.header("Enter Your Project Description")
# query = st.text_area("Enter a project description for flowchart generation:")

# # Button to generate flowchart
# if st.button("Generate Flowchart"):
#     if uploaded_file is not None and query:
#         st.text("Extracting text from the uploaded file...")
        
#         # Extract text from the uploaded DOCX file
#         knowledge_base_text = extract_text_from_single_word_file(uploaded_file)
        
#         st.text("Creating embeddings...")
        
#         # Create embeddings for knowledge base text
#         embed = openai.Embedding.create(
#             input=knowledge_base_text,
#             engine="text-embedding-ada-002"
#         )
#         time.sleep(10)
        
#         st.text("Generating flowchart description...")
        
#         # Generate a response based on the prompt with knowledge base text as context
#         response = openai.Completion.create(
#             model="text-davinci-002",
#             prompt=f"You are a Data Architect tasked with designing an optimized data science project pipeline, including infrastructure requirements. Describe a high-level flowchart that outlines the step-by-step process for building the project from scratch. Your flowchart should be based on the following information: {query}. Knowledge Base: {knowledge_base_text}",
#             temperature=0,
#             max_tokens=500
#         )
        
#         flowchart_description = response.choices[0].text
        
#         # Display the generated flowchart description
#         st.header("Generated Flowchart Description:")
#         st.text(flowchart_description)
        
#         # Create the interactive flowchart and display it
#         flowchart_image = create_interactive_flowchart(flowchart_description)
#         st.header("Interactive Flowchart:")
#         st.image(flowchart_image)
#     else:
#         st.error("Please upload a DOCX file and enter a project description.")

# st.sidebar.text("Note: Ensure you have entered your OpenAI API key.")


# import streamlit as st
# import os
# from docx import Document
# import re
# import openai
# import time
# from PIL import Image
# import subprocess

# # Function to extract text from a single Word file
# def extract_text_from_single_word_file(file):
#     doc = Document(file)
#     text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
#     clean_text = re.sub(r'[^\w\s]', '', text)

#     if len(clean_text) > 409050:
#         clean_text = clean_text[:400000]

#     extracted_text = clean_text[:40956]
#     return extracted_text

# # Function to create a flowchart from the generated description
# def create_flowchart(description):
#     dot_file_path = 'flowchart.dot'
#     png_file_path = 'flowchart.png'
    
#     # Create a DOT file for the flowchart
#     with open(dot_file_path, 'w') as dot_file:
#         dot_file.write('digraph G {\n')
        
#         # Split the description into lines
#         lines = description.strip().split('\n')
        
#         for line in lines:
#             parts = line.strip().split('->')
#             if len(parts) == 2:
#                 src, dest = parts
#                 dot_file.write(f'  "{src.strip()}" -> "{dest.strip()}"\n')
        
#         dot_file.write('}\n')
    
#     # Use Graphviz to generate the PNG image from the DOT file
#     subprocess.run(['dot', '-Tpng', dot_file_path, '-o', png_file_path])
    
#     return png_file_path

# # Initialize OpenAI API
# openai.api_key = 'sk-'  # Replace with your actual OpenAI API key

# # Streamlit UI
# st.title("Flowchart Generation from Text")

# # File upload
# st.header("Upload DOCX File")
# uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"])

# # Prompt input
# st.header("Enter Your Project Description")
# query = st.text_area("Enter a project description for flowchart generation:")

# # Button to generate flowchart
# if st.button("Generate Flowchart"):
#     if uploaded_file is not None and query:
#         st.text("Extracting text from the uploaded file...")

#         # Extract text from the uploaded DOCX file
#         knowledge_base_text = extract_text_from_single_word_file(uploaded_file)

#         st.text("Generating flowchart description...")

#         # Generate a response based on the prompt with knowledge base text as context
#         response = openai.Completion.create(
#             model="text-davinci-002",
#             prompt=f"Create a flowchart based on the following description: {query}. Knowledge Base: {knowledge_base_text}",
#             temperature=0,
#             max_tokens=500
#         )

#         flowchart_description = response.choices[0].text

#         # Display the generated flowchart description
#         st.header("Generated Flowchart Description:")
#         st.text(flowchart_description)

#         # Create the flowchart and display it
#         flowchart_image = create_flowchart(flowchart_description)
#         st.header("Flowchart:")
#         st.image(flowchart_image)
#     else:
#         st.error("Please upload a DOCX file and enter a project description.")

# st.sidebar.text("Note: Ensure you have entered your OpenAI API key.")
###################### Manual Flow  ###############################
# import streamlit as st
# import openai
# from graphviz import Digraph

# # Set your OpenAI API key here
# openai.api_key = "sk-xyz"

# # Streamlit UI
# st.title("Workflow Generator")

# # User input
# user_input = st.text_area("Enter your project/product description:")
# if st.button("Generate Workflow"):
#     if not user_input:
#         st.warning("Please enter a description.")
#     else:
#         # Generate workflow using OpenAI GPT-3
#         response = openai.Completion.create(
#             engine="text-davinci-002",  # Use the appropriate engine for your needs
#             prompt= 'You are a Data Architect .You are responsible for creating a work flow based on project requirements when a user gives  {user_input} for generating flow chart you must analyse and generate detailed flow diagram step by step. The first step is to understand the project requirements and gather the necessary data. The next step is to design the data architecture and build the data pipelines. The next step is to deploy the data architecture and pipelines and you have to generate flowchart which has the optimised solution for the complete pipelines to build a product from scratch by taking in the infrastructure requirements of a project related to data science .Please make sure that you are giving the tools name inside the shapes ',
#             max_tokens=100,  # Adjust as needed
#             n=1,
#             stop=None,
#             temperature=0.7,
#         )
#         workflow_text = response.choices[0].text

#         # Create a simple diagram using Graphviz
#         dot = Digraph(comment='Workflow Diagram')
#         dot.node('A', 'Start', shape='ellipse', color='green')
#         dot.node('B', 'Data Collection', shape='box', color='blue')
#         dot.node('C', 'Data Preprocessing', shape='box', color='blue')
#         dot.node('D', 'Feature Engineering', shape='box', color='blue')
#         dot.node('E', 'Model Training', shape='box', color='blue')
#         dot.node('F', 'Model Evaluation', shape='box', color='blue')
#         dot.node('G', 'Model Deployment', shape='box', color='blue')
#         dot.node('H', 'End', shape='ellipse', color='red')

#         dot.edges(['AB', 'BC', 'CD', 'DE', 'EF', 'FG', 'GH'])

#         with dot.subgraph() as s:
#             s.attr(rank='same')
#             s.node('C', 'Data Preprocessing', shape='box', color='blue')
#             s.node('D', 'Feature Engineering', shape='box', color='blue')
#         with dot.subgraph() as s:
#             s.attr(rank='same')
#             s.node('E', 'Model Training', shape='box', color='blue')
#             s.node('F', 'Model Evaluation', shape='box', color='blue')

#         # Display the generated diagram
#         st.graphviz_chart(dot.source)


# import streamlit as st
# import openai
# from graphviz import Digraph

# # Set your OpenAI API key here
# openai.api_key = "sk-"

# # Streamlit UI
# st.title("Workflow Generator")

# # User input
# user_input = st.text_area("Enter your project/product description:")
# if st.button("Generate Workflow"):
#     if not user_input:
#         st.warning("Please enter a description.")
#     else:
#         # Generate workflow using OpenAI GPT-3
#         response = openai.Completion.create(
#             engine="text-davinci-002",  # Use the appropriate engine for your needs
#             prompt=f'You are a Data Architect. You are responsible for creating a workflow based on project requirements when a user gives "{user_input}" for generating a flowchart. You must analyze and generate a detailed flow diagram step by step. Please make sure that you include the tools\' names inside the shapes.',
#             max_tokens=150,  # Adjust as needed
#             n=1,
#             stop=None,
#             temperature=0.7,
#         )
#         workflow_text = response.choices[0].text

#         # Parse the response into steps
#         steps = workflow_text.split('\n')

#         # Create a simple diagram using Graphviz
#         dot = Digraph(comment='Workflow Diagram')

#         # Define shapes and colors for nodes
#         shapes = ['ellipse', 'box', 'diamond', 'triangle', 'parallelogram', 'hexagon', 'octagon']
#         colors = ['green', 'blue', 'red', 'purple', 'orange', 'pink', 'yellow']

#         for i, step in enumerate(steps):
#             # Choose a shape and color for each step
#             shape = shapes[i % len(shapes)]
#             color = colors[i % len(colors)]
#             dot.node(str(i), step, shape=shape, color=color)

#         # Add edges between the steps
#         for i in range(len(steps) - 1):
#             dot.edge(str(i), str(i + 1))

#         # Display the generated diagram
#         st.graphviz_chart(dot.source)


# import re
# import streamlit as st
# import openai
# from graphviz import Digraph

# # Set your OpenAI API key here
# openai.api_key = "sk-"

# # Streamlit UI
# st.set_page_config(
#     page_title="Workflow Generator",
#     page_icon="🚀",
#     layout="wide",
#     initial_sidebar_state="expanded",
# )

# # Custom CSS for styling
# custom_css = """
# <style>
#     /* Custom styles */
#     body {
#         font-family: Arial, sans-serif;
#         background-color: #f5f5f5;
#         margin: 0;
#         padding: 0;
#     }
#     .stApp {
#         max-width: 100%;
#         padding: 20px;
#         background-color: #fff;
#         box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
#         border-radius: 10px;
#     }
#     .stButton {
#         background-color: #007BFF;
#         color: #fff;
#         border: none;
#         border-radius: 5px;
#         padding: 10px 20px;
#         cursor: pointer;
#         font-size: 16px;
#     }
#     .stButton:hover {
#         background-color: #0056b3;
#     }
#     .stSidebar {
#         background-color: #f5f5f5;
#         padding: 20px;
#         border-right: 1px solid #ddd;
#     }
# </style>
# """

# st.markdown(custom_css, unsafe_allow_html=True)
# st.title("Workflow Generator")
# st.sidebar.header("Settings")

# # User input
# user_input = st.text_area("Enter your project/product description:", height=150)
# generate_button = st.button("Generate Workflow")

# if not user_input:
#     st.warning("Please enter a description.")
# elif generate_button:
#     # Generate a technical workflow using OpenAI GPT-3 with a prompt emphasizing technical details
#     prompt = f'You are a Data Architect tasked with designing a technical workflow. The project description provided is as follows:\n\n{user_input}\n\nCommence by elucidating the fundamental technical steps and essential components required to establish the workflow. Please delve into technical specifics, including tools, technologies, and methodologies.'

#     response = openai.Completion.create(
#         engine="text-davinci-002",
#         prompt=prompt,
#         max_tokens=300,
#         n=1,
#         stop=None,
#         temperature=0.7,
#     )
#     workflow_text = response.choices[0].text

#     # Extract and remove numbering from the response
#     steps = [re.sub(r"^\d+\.\s*", "", step.strip()) for step in workflow_text.split('\n') if step.strip()]

#     dot = Digraph(comment='Technical Workflow Diagram')
#     dot.graph_attr['size'] = '10,10'

#     # Customize the appearance of nodes and edges
#     dot.attr('node', shape='rectangle', style='filled', fillcolor='#E4F2F2', fontname='Arial', fontsize='16', color='#3D7E9A')
#     dot.attr('edge', color='#007BFF', penwidth='2', dir='none', arrowhead='none', fontname='Arial')

#     for i, step in enumerate(steps):
#         dot.node(f'step_{i}', step)

#     # Create diagonal edges between the steps
#     for i in range(len(steps) - 1):
#         dot.edge(f'step_{i}', f'step_{i + 1}', dir='none')

#     st.header("Workflow Diagram")
#     st.graphviz_chart(dot.source)

# # Add a footer with attribution
# st.sidebar.markdown("---")
# st.sidebar.markdown("*Workflow Generator* by Your Name")
# ---------------------------------------


# import re
# import streamlit as st
# import openai
# from graphviz import Digraph

# # Set your OpenAI API key here
# openai.api_key = "sk-"

# # Custom CSS for styling
# custom_css = """
# <style>
#     /* Custom styles */
#     body {
#         font-family: Arial, sans-serif;
#         background-color: #f5f5f5;
#         margin: 0;
#         padding: 0;
#     }
#     .stApp {
#         max-width: 100%;
#         padding: 20px;
#         background-color: #fff;
#         box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
#         border-radius: 10px;
#     }
#     .stButton {
#         background-color: #007BFF;
        
#         border: none;
#         border-radius: 5px;
#         padding: 10px 20px;
#         cursor: pointer;
#         font-size: 16px;
#     }
#     .stButton:hover {
#         background-color: #0056b3;
#     }
# </style>
# """

# st.markdown(custom_css, unsafe_allow_html=True)
# st.title("Workflow Generator")

# # User input
# user_input = st.text_area("Enter your project/product description:", height=150)
# generate_button = st.button("Generate Workflow")

# if not user_input:
#     st.warning("Please enter a description.")
# elif generate_button:
#     # Generate a technical workflow using OpenAI GPT-3 with a prompt emphasizing technical details
#     prompt = f'You are a Data Architect tasked with designing a technical workflow. The project description provided is as follows:\n\n{user_input}\n\nCommence by elucidating the fundamental technical steps and essential components to establish the workflow. Please delve into technical specifics, including tools, technologies, and methodologies.Make sure to suggest tools and tecnologoies that can be used along with the solution'

#     response = openai.Completion.create(
#         engine="text-davinci-002",
#         prompt=prompt,
#         max_tokens=300,
#         n=1,
#         stop=None,
#         temperature=0.0,
#     )
  
#     workflow_text = response.choices[0].text

#     # Extract and remove numbering from the response
#     steps = [re.sub(r"^\d+\.\s*", "", step.strip()) for step in workflow_text.split('\n') if step.strip()]

#     dot = Digraph(comment='Technical Workflow Diagram')
#     dot.graph_attr['size'] = '10,10'

#     # Customize the appearance of nodes and edges
#     dot.attr('node', shape='rectangle', style='filled', fillcolor='#E4F2F2', fontname='Arial', fontsize='16', color='#3D7E9A')
#     dot.attr('edge', color='#007BFF', penwidth='2', dir='none', arrowhead='none', fontname='Arial')

#     for i, step in enumerate(steps):
#         dot.node(f'step_{i}', step)

#     # Create diagonal edges between the steps
#     for i in range(len(steps) - 1):
#         dot.edge(f'step_{i}', f'step_{i + 1}', dir='none')

#     st.header("Workflow Diagram")
#     st.graphviz_chart(dot.source)



import streamlit as st
from docx import Document
from graphviz import Digraph
import openai

# Set your OpenAI API key here
openai.api_key = "sk-"

# Streamlit UI
st.set_page_config(
    page_title="Flowchart Generator",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Custom CSS for styling
custom_css = """
<style>
    /* Custom styles */
    body {
        font-family: Arial, sans-serif;
        background-color: #333; /* Dark background */
        color: #fff; /* White text */
        margin: 0;
        padding: 0;
    }
    .stApp {
        max-width: 800px;
        margin: 0 auto;
        padding: 20px;
        background-color: #444; /* Slightly lighter than the background */
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
        border-radius: 10px;
    }
    .stButton {
        background-color: #007BFF;
        color: #fff;
        border: none;
        border-radius: 5px;
        padding: 10px 20px;
        cursor: pointer;
        font-size: 16px;
    }
    .stButton:hover {
        background-color: #0056b3;
    }
    .stHeader {
        font-size: 24px;
        margin-bottom: 20px;
    }
    .stText {
        font-size: 16px;
    }
</style>
"""

st.markdown(custom_css, unsafe_allow_html=True)
st.title("Flowchart Generator")

# File upload
st.header("Upload a DOCX file")
uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"])

# User input for description
st.header("Enter a project/product description:")
user_description = st.text_area("Description", height=150)

# Button to generate the flowchart
generate_button = st.button("Generate Flowchart")

if not uploaded_file:
    st.warning("Please upload a DOCX file.")
elif not user_description:
    st.warning("Please provide a description.")
elif generate_button:
    # Extract text from the uploaded DOCX file
    doc = Document(uploaded_file)
    doc_text = ' '.join([paragraph.text for paragraph in doc.paragraphs])

    # Combine user description and doc text
    combined_text = f"{user_description.strip()}\n\n{doc_text}"

    # Generate a refined flowchart description using GPT-3
    response = openai.Completion.create(
        engine="text-davinci-002",
        prompt=f"Generate a refined flowchart based on the following description:\n\n{combined_text}\n\n",
        max_tokens=300,
    )
    refined_description = response.choices[0].text

    steps = [step.strip() for step in refined_description.split('\n') if step.strip()]
    dot = Digraph(comment='Refined Flowchart')
    dot.graph_attr['size'] = '10,10'

    for i, step in enumerate(steps):
        dot.node(f'step_{i}', step)

    for i in range(len(steps) - 1):
        dot.edge(f'step_{i}', f'step_{i + 1}')

    st.header("Generated Refined Flowchart:")
    st.graphviz_chart(dot.source)
